"""Claude Haiku prompts for Phase 3 assignment workspace."""

import json
import logging
import httpx
from pathlib import Path

from ..config import settings

log = logging.getLogger(__name__)

HAIKU = "claude-haiku-4-5-20251001"


def _call_claude(system: str, user: str, max_tokens: int = 2048) -> str:
    response = httpx.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": settings.anthropic_api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": HAIKU,
            "max_tokens": max_tokens,
            "system": system,
            "messages": [{"role": "user", "content": user}],
        },
        timeout=180,
    )
    response.raise_for_status()
    return response.json()["content"][0]["text"]


def _call_claude_vision(system: str, user_text: str, image_paths: list[str], max_tokens: int = 2048) -> str:
    import base64
    content = []
    ext_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
                ".gif": "image/gif", ".webp": "image/webp"}
    for p in image_paths:
        path = Path(p)
        if not path.exists():
            continue
        media_type = ext_map.get(path.suffix.lower(), "image/png")
        data = base64.standard_b64encode(path.read_bytes()).decode()
        content.append({"type": "image", "source": {"type": "base64", "media_type": media_type, "data": data}})
    content.append({"type": "text", "text": user_text})

    response = httpx.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": settings.anthropic_api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": HAIKU,
            "max_tokens": max_tokens,
            "system": system,
            "messages": [{"role": "user", "content": content}],
        },
        timeout=180,
    )
    response.raise_for_status()
    return response.json()["content"][0]["text"]


def _extra(instructions: str) -> str:
    if instructions.strip():
        return f"\n\nAdditional instructions from the student:\n{instructions.strip()}"
    return ""


def get_hint(question: str, solution_so_far: str, extra_instructions: str = "") -> str:
    system = """You are a helpful programming tutor. The student is working on a CS assignment.
Give a concise, specific hint (2-3 sentences max) about what they should do next.
Do NOT give away the complete answer. Focus on the immediate next concept or step to apply."""

    user = f"""Assignment Question:
{question}

Student's solution so far:
{solution_so_far.strip() if solution_so_far.strip() else "(nothing written yet)"}

What should the student focus on next?{_extra(extra_instructions)}"""

    return _call_claude(system, user, max_tokens=300)


def _strip_code_fences(text: str) -> str:
    """Remove markdown code fences from AI responses."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        # drop first line (```lang) and last line (```)
        end = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
        text = "\n".join(lines[1:end]).strip()
    return text


def complete_solution(question: str, roll_number: str, extra_instructions: str = "") -> str:
    system = f"""You are an expert CS student completing an assignment.
Your VU roll number is {roll_number} — use it wherever the assignment asks you to perform operations on your roll number (e.g. use digits of roll number as input values, perform calculations on it, etc.).
Write a complete, correct solution to the assignment question.
Write ONLY the answer content — no name, roll number headers, file metadata, or submission boilerplate.
For code assignments: output raw code only, no markdown fences, no explanation."""

    raw = _call_claude(system, f"Assignment:\n\n{question}{_extra(extra_instructions)}", max_tokens=4096)
    return _strip_code_fences(raw)


def format_for_upload(
    question: str,
    roll_number: str,
    student_name: str,
    course_name: str,
    extra_instructions: str = "",
    image_paths: list[str] = [],
) -> dict:
    """
    Returns dict with:
      filename: str  (without extension)
      extension: str
      code: str      (Python 3 code; `output_path` and `solution_text` are pre-injected)
    """
    default_filename = f"{roll_number}_{course_name}"

    system = f"""You are a Python developer. Write Python 3 code that saves a student's assignment solution to a file.

Two variables are already defined when your code runs — do NOT redefine them:
- `output_path`: the file path where you must save the file (a string)
- `solution_text`: the student's complete solution text (a string)

Your code must:
1. Determine the correct file format from the assignment question (default: .docx)
2. Create the document at `output_path` using `solution_text` as the content
3. For .docx files: use python-docx (`from docx import Document`)
4. For code files (.cpp, .c, .py, .java, .cs, etc.) and .txt: use `open(output_path, 'w', encoding='utf-8')`
5. For code files: strip any markdown fences from solution_text before writing. Use this snippet:
   `import re; code = re.sub(r'^```[\\w]*\\n?', '', solution_text.strip(), flags=re.MULTILINE); code = re.sub(r'^```$', '', code, flags=re.MULTILINE); code = code.strip()`
6. ONLY include student name or roll number in the document if the assignment question EXPLICITLY instructs students to write their name or ID in the file. If not mentioned, omit them entirely.
7. Write content using the native conventions of the target file format. For example, use real Word headings, paragraphs, and tables in a .docx file — not markdown syntax. Use proper code structure in a .cpp/.py file. Use plain readable text in a .txt file. Never apply formatting idioms from one format to another (e.g., no `# Heading` or `| col |` in a Word document).
8. The variable `image_paths` is a pre-defined Python list of file paths to screenshots provided by the student. For document formats (.docx): embed each image at a contextually appropriate place using `from docx.shared import Inches` and `doc.add_picture(p, width=Inches(5.5))`. For code files (.cpp, .c, .py, .java, etc.): images CANNOT be embedded in source code — ignore `image_paths` entirely and just write the solution code to the file. The screenshots were provided for your visual context only.

Example for a .docx output:
```
from docx import Document
doc = Document()
doc.add_heading('{course_name} Assignment', 0)
doc.add_paragraph(solution_text)
doc.save(output_path)
```

Return ONLY valid JSON with exactly these three keys — no markdown fences, no explanation:
{{"filename": "name_without_extension", "extension": "docx", "code": "python 3 code with \\n for newlines"}}

The default filename is: {default_filename}"""

    images_note = f"\n\n{len(image_paths)} screenshot(s) are attached — embed them appropriately in the document." if image_paths else ""

    user = f"""Assignment Question:
{question}

Student Name: {student_name}
Roll Number: {roll_number}
Course: {course_name}

What file format does this assignment require? Write Python code to create that file using `solution_text`.{_extra(extra_instructions)}{images_note}"""

    if image_paths:
        raw = _call_claude_vision(system, user, image_paths, max_tokens=2048)
    else:
        raw = _call_claude(system, user, max_tokens=2048)

    # Strip markdown fences if Claude wraps the JSON
    if "```json" in raw:
        raw = raw.split("```json")[1].split("```")[0].strip()
    elif "```" in raw:
        raw = raw.split("```")[1].split("```")[0].strip()

    return json.loads(raw.strip())
