from datetime import datetime, timezone
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session

from ..db import get_db
from ..models import Item
from ..schemas import ItemOut

router = APIRouter(prefix="/api/items", tags=["items"])


def _item_out(item: Item) -> ItemOut:
    return ItemOut(
        id=item.id,
        course_id=item.course_id,
        course_code=item.course.code,
        course_title=item.course.title,
        kind=item.kind,
        title=item.title,
        lesson=item.lesson,
        total_marks=item.total_marks,
        status=item.status,
        opens_at=item.opens_at,
        due_at=item.due_at,
        file_url=item.file_url,
        completed_at=item.completed_at,
        first_seen_at=item.first_seen_at,
        last_seen_at=item.last_seen_at,
    )


@router.get("", response_model=list[ItemOut])
def list_items(
    due_within_days: int = 14,
    kind: str = None,
    db: Session = Depends(get_db),
):
    from datetime import timedelta
    from sqlalchemy import or_
    now = datetime.now(timezone.utc)
    cutoff = now + timedelta(days=due_within_days)
    q = db.query(Item).filter(
        Item.completed_at.is_(None),
    ).filter(
        or_(
            # upcoming items within window
            (Item.due_at >= now) & (Item.due_at <= cutoff),
            # past-due but LMS still shows open
            (Item.due_at < now) & (Item.status == "Open"),
        )
    )
    if kind:
        q = q.filter(Item.kind == kind)
    items = q.order_by(Item.due_at).all()
    return [_item_out(i) for i in items]


@router.get("/{item_id}", response_model=ItemOut)
def get_item(item_id: int, db: Session = Depends(get_db)):
    item = db.query(Item).get(item_id)
    if not item:
        raise HTTPException(404, "Item not found")
    return _item_out(item)


async def _fetch_and_cache_file(item: Item) -> tuple[bytes, str]:
    """Download assignment file via Playwright and cache it locally. Returns (bytes, filename)."""
    from ..scraper import VULMSScraper, CourseDTO
    from ..config import settings

    course = item.course
    course_dto = CourseDTO(
        index=course.lms_index,
        ctl_id=course.ctl_id,
        postback_target=course.postback_target,
        code=course.code,
        title=course.title,
    )

    scraper = VULMSScraper()
    try:
        await scraper.start(headless=True)
        ok = await scraper.ensure_logged_in()
        if not ok:
            raise HTTPException(503, "Could not authenticate with LMS")

        data, filename = await scraper.download_assignment_file(course_dto, item.lms_index)

        # Cache the raw file
        cache_path = settings.files_dir / f"{item.id}_{filename}"
        cache_path.write_bytes(data)

        # Also write a pointer file so we can find it without knowing the filename
        ptr = settings.files_dir / f"{item.id}.ptr"
        ptr.write_text(filename)

        return data, filename
    finally:
        await scraper.stop()


def _get_cached_file(item_id: int) -> tuple[bytes, str] | None:
    """Return (bytes, filename) from cache, or None if not cached."""
    from ..config import settings

    ptr = settings.files_dir / f"{item_id}.ptr"
    if not ptr.exists():
        return None
    filename = ptr.read_text().strip()
    cache_path = settings.files_dir / f"{item_id}_{filename}"
    if not cache_path.exists():
        return None
    return cache_path.read_bytes(), filename


@router.get("/{item_id}/text")
async def get_text(item_id: int, db: Session = Depends(get_db)):
    """Extract plain text from the assignment file for AI context. Cached after first fetch."""
    from ..config import settings

    item = db.query(Item).get(item_id)
    if not item:
        raise HTTPException(404, "Item not found")

    text_cache = settings.files_dir / f"{item_id}.txt"
    if text_cache.exists():
        return {"text": text_cache.read_text()}

    cached = _get_cached_file(item_id)
    if cached:
        data, filename = cached
    else:
        try:
            data, filename = await _fetch_and_cache_file(item)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Download failed: {e}")

    file_path = settings.files_dir / f"{item_id}_{filename}"
    try:
        text = _extract_text(file_path)
        text_cache.write_text(text)
        return {"text": text}
    except Exception as e:
        raise HTTPException(500, f"Text extraction failed: {e}")


@router.get("/{item_id}/file")
async def download_file(item_id: int, db: Session = Depends(get_db)):
    """Download raw assignment file (.docx). Cached after first fetch."""
    from fastapi.responses import StreamingResponse
    import io

    item = db.query(Item).get(item_id)
    if not item:
        raise HTTPException(404, "Item not found")
    if item.kind != "assignment":
        raise HTTPException(400, "Only assignment files supported")

    cached = _get_cached_file(item_id)
    if cached:
        data, filename = cached
    else:
        try:
            data, filename = await _fetch_and_cache_file(item)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Download failed: {e}")

    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _extract_text(file_path) -> str:
    """Extract plain text from a .doc or .docx file."""
    import subprocess
    import tempfile
    import shutil
    from pathlib import Path
    from docx import Document

    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # .doc — convert to .docx via LibreOffice first
    with tempfile.TemporaryDirectory() as tmpdir:
        shutil.copy(file_path, tmpdir)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", tmpdir, str(file_path)],
            capture_output=True, timeout=60,
        )
        converted = Path(tmpdir) / (file_path.stem + ".docx")
        if converted.exists():
            doc = Document(str(converted))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise RuntimeError(f"Could not convert {file_path.name} to text")


def _word_to_html(file_path) -> str:
    """Convert a .doc or .docx file to HTML body content."""
    import mammoth
    import io
    import subprocess
    import tempfile
    import shutil
    from pathlib import Path

    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        result = mammoth.convert_to_html(io.BytesIO(file_path.read_bytes()))
        return result.value

    # .doc or unknown — convert to .docx via LibreOffice first, then mammoth
    with tempfile.TemporaryDirectory() as tmpdir:
        shutil.copy(file_path, tmpdir)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", tmpdir, str(file_path)],
            capture_output=True, timeout=60,
        )
        converted = Path(tmpdir) / (file_path.stem + ".docx")
        if converted.exists():
            result = mammoth.convert_to_html(io.BytesIO(converted.read_bytes()))
            return result.value

    raise RuntimeError(f"Could not convert {file_path.name} to HTML")


PAGE_CSS = """
  body {
    font-family: Georgia, 'Times New Roman', serif;
    font-size: 14px;
    line-height: 1.8;
    color: #1a1a1a;
    max-width: 780px;
    margin: 0 auto;
    padding: 32px 24px 64px;
    background: #fff;
  }
  h1, h2, h3, h4 { font-weight: 700; margin: 1.2em 0 0.4em; color: #111; }
  h1 { font-size: 1.5em; }
  h2 { font-size: 1.25em; }
  h3 { font-size: 1.1em; }
  p { margin: 0.6em 0; }
  table { border-collapse: collapse; width: 100%; margin: 1em 0; }
  td, th { border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; }
  th { background: #f3f4f6; font-weight: 600; }
  ul, ol { padding-left: 1.5em; margin: 0.6em 0; }
  li { margin: 0.3em 0; }
  strong { font-weight: 700; }
  em { font-style: italic; }
  img { max-width: 100%; height: auto; }
  pre, code { font-family: monospace; background: #f3f4f6; padding: 2px 6px; border-radius: 3px; font-size: 13px; }
"""


@router.get("/{item_id}/file/view")
async def view_file(item_id: int, db: Session = Depends(get_db)):
    """Convert assignment Word file to HTML for inline viewing. Result is cached."""
    from fastapi.responses import HTMLResponse
    from ..config import settings

    item = db.query(Item).get(item_id)
    if not item:
        raise HTTPException(404, "Item not found")
    if item.kind != "assignment":
        raise HTTPException(400, "Only assignment files supported")

    # Serve from HTML cache if available
    html_cache = settings.files_dir / f"{item_id}.html"
    if html_cache.exists():
        return HTMLResponse(content=html_cache.read_text())

    # Ensure raw file is downloaded
    cached = _get_cached_file(item_id)
    if cached:
        _, filename = cached
    else:
        try:
            _, filename = await _fetch_and_cache_file(item)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Download failed: {e}")

    # Find the cached file path
    file_path = settings.files_dir / f"{item_id}_{filename}"

    try:
        body_html = _word_to_html(file_path)
    except Exception as e:
        raise HTTPException(500, f"Conversion failed: {e}")

    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>{PAGE_CSS}</style></head><body>{body_html}</body></html>"

    # Cache the HTML so future views are instant
    html_cache.write_text(html)

    return HTMLResponse(content=html)
